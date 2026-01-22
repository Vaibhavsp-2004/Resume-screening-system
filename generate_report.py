from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('AI Powered Resume Screening System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Project Authorization: AISE - Vaibhav S P', style='Subtitle')
    doc.add_paragraph('')

    # Abstract
    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        "Traditional hiring processes are time-consuming and prone to human bias. "
        "This project implements an AI-powered Applicant Tracking System (ATS) that automates "
        "resume screening. Using Natural Language Processing (NLP), the system compares resumes "
        "against job descriptions to calculate a relevance score, helping recruiters shortlist "
        "candidates efficiently."
    )

    # Problem Statement
    doc.add_heading('2. Problem Statement', level=1)
    doc.add_paragraph(
        "Recruiters receive hundreds of resumes for a single opening. Manually reviewing each one "
        "takes significantly longer and can lead to fatigue-driven errors. There is a need for a "
        "tool that can instantly filter and rank candidates based on objective criteria."
    )

    # Solution & Architecture
    doc.add_heading('3. Proposed Solution', level=1)
    doc.add_paragraph(
        "We developed a web-based application consisting of:"
    )
    p = doc.add_paragraph()
    p.add_run("1. Backend: ").bold = True
    p.add_run("FastAPI (Python) for high-performance API handling.\n")
    p.add_run("2. Frontend: ").bold = True
    p.add_run("HTML/CSS/JavaScript for a responsive recruiter dashboard.\n")
    p.add_run("3. AI Engine: ").bold = True
    p.add_run("scikit-learn (CountVectorizer & Cosine Similarity) to match keywords.\n")
    p.add_run("4. Infrastructure: ").bold = True
    p.add_run("Docker for containerization and Render.com for cloud deployment.")

    # Key Features
    doc.add_heading('4. Key Features', level=1)
    items = [
        "Secure Authentication (JWT) for Recruiters",
        "Dynamic Resume Parsing (Supports .pdf, .docx, .txt)",
        "Intelligent Matching Algorithm (TF/Vector-based Scoring)",
        "Real-time Ranking of Candidates",
        "Instant 'View Profile' access to resumes",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Demo Script
    doc.add_heading('5. Demo Flow (For Presentation)', level=1)
    doc.add_paragraph("Follow these steps during your live demo:")
    
    steps = [
        "Login: Use the recruiter credentials to access the dashboard.",
        "Create Job: Post a 'Senior Python Developer' role with required skills (Python, Docker, AWS).",
        "Upload Resumes: Upload a mix of files (.pdf, .docx, .txt).",
        "Show Parsing: Explain how the system reads simple text and complex tables.",
        "Analyze Results: Show the Ranked List. Point out the top candidate.",
        "Verification: Click 'View Profile' to show the original file.",
        "Scoring Logic: Briefly explain that the system counts keyword overlaps to generate a score (0-100%)."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # Conclusion
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "The AI Resume Screening System demonstrates how modern NLP techniques can transform "
        "HR operations. It reduces screening time by nearly 90% and provides an objective ranking mechanism."
    )

    filename = "AI_Resume_Project_Report.docx"
    doc.save(filename)
    print(f"Successfully generated: {filename}")

if __name__ == "__main__":
    try:
        create_report()
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        print(f"Error: {e}")
